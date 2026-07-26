from __future__ import annotations

import os
import re
import uuid
import zipfile
from collections import OrderedDict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from PIL import Image

from .content_blocks import content_block_rows, content_blocks_for_problem
from .formula_math import FormulaValidationError, latex_to_omml


@dataclass(frozen=True)
class DocxExportResult:
    path: Path
    filename: str


_LATEX_TOKEN = re.compile(
    r"\d+(?:\.\d+)?|[A-Za-z]|\\(?:times|div|le|ge)|[+\-=:\/]"
)
_UNSUPPORTED_MATH = re.compile(r"[\u221a\u2211\u222b\u2248\u2260]|\\[A-Za-z]+")
_OPERATOR_TEXT = {
    r"\times": "\u00d7",
    r"\div": "\u00f7",
    r"\le": "\u2264",
    r"\ge": "\u2265",
}


def docx_exportable(problem: dict[str, Any]) -> bool:
    category = str(problem.get("category") or "")
    category_key = str(problem.get("category_key") or "")
    return bool(
        problem.get("selected_artifact")
        and problem.get("category_group")
        and category_key
        and category == category_key
        and problem.get("ocr_text")
        and (
            problem.get("status") == "ready"
            or problem.get("review_status") == "accepted"
        )
    )


def _group_problems(
    problems: list[dict[str, Any]],
) -> OrderedDict[str, OrderedDict[str, list[dict[str, Any]]]]:
    grouped: OrderedDict[str, OrderedDict[str, list[dict[str, Any]]]] = OrderedDict()
    for problem in problems:
        group = str(problem.get("category_group") or "\u672a\u5206\u7c7b")
        category = str(
            problem.get("category_key")
            or problem.get("category")
            or "\u672a\u5206\u7c7b"
        )
        grouped.setdefault(group, OrderedDict()).setdefault(category, []).append(problem)
    return grouped


def _problem_title(problem: dict[str, Any], number: int) -> str:
    structured = (problem.get("metrics") or {}).get("structured_problem", {})
    title = structured.get("title") if isinstance(structured, dict) else ""
    if title and title != "\u3010\u9898\u76ee\u3011":
        return str(title)
    text = str(problem.get("ocr_text") or "").strip()
    first_line = text.splitlines()[0].strip() if text else ""
    match = re.match(
        r"^((?:\u3010|\[).{1,24}?(?:\u3011|\]))",
        first_line,
    )
    if match:
        return match.group(1)
    return f"\u9898\u76ee {number}"


def _set_run_font(run: Any, size: float = 11) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _configure_document(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    for style_name, size in (
        ("Normal", 11),
        ("Title", 20),
        ("Heading 1", 16),
        ("Heading 2", 13),
        ("Heading 3", 11),
    ):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style._element.get_or_add_rPr().rFonts.set(
            qn("w:eastAsia"), "Microsoft YaHei"
        )
def _add_toc(document: Any) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "\u5728 Word \u4e2d\u66f4\u65b0\u57df\u540e\u663e\u793a\u76ee\u5f55"
    separate_run = OxmlElement("w:r")
    separate_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, separate_run, end])


def _formula_text(latex: str) -> str:
    compact = re.sub(r"\s+", "", latex)
    tokens = _LATEX_TOKEN.findall(compact)
    if not tokens or "".join(tokens) != compact:
        raise ValueError(f"\u4e0d\u652f\u6301\u7684 LaTeX \u516c\u5f0f\uff1a{latex}")
    return "".join(_OPERATOR_TEXT.get(token, token) for token in tokens)


def _append_omml(paragraph: Any, latex: str) -> None:
    try:
        xml = latex_to_omml(latex)
    except FormulaValidationError:
        math = OxmlElement("m:oMath")
        math_run = OxmlElement("m:r")
        math_text = OxmlElement("m:t")
        math_text.text = _formula_text(latex)
        math_run.append(math_text)
        math.append(math_run)
        paragraph._p.append(math)
        return
    xml = xml.replace(
        "<m:oMath>",
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/'
        'officeDocument/2006/math">',
        1,
    )
    paragraph._p.append(parse_xml(xml))


def _resource_path(problem: dict[str, Any], asset: str) -> Path:
    if Path(asset).name != asset:
        raise ValueError("\u9898\u76ee\u8d44\u6e90\u8def\u5f84\u4e0d\u5b89\u5168")
    artifact_dir = Path(str(problem.get("selected_artifact") or "")).resolve().parent
    unresolved = artifact_dir / asset
    candidate = unresolved.resolve()
    if unresolved.is_symlink() or candidate.parent != artifact_dir or not candidate.is_file():
        raise ValueError(
            f"\u9898\u76ee {problem.get('filename') or problem.get('id')} "
            f"\u7f3a\u5c11\u8d44\u6e90 {asset}"
        )
    return candidate


def _add_image(document: Any, path: Path) -> None:
    with Image.open(path) as image:
        width, height = image.size
    max_width = 150.0
    max_height = 70.0
    ratio = width / max(1, height)
    draw_width = min(max_width, max_height * ratio)
    draw_height = draw_width / ratio
    if draw_height > max_height:
        draw_height = max_height
        draw_width = draw_height * ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture(
        str(path), width=Mm(draw_width), height=Mm(draw_height)
    )


def _add_inline_image(paragraph: Any, path: Path) -> None:
    with Image.open(path) as image:
        width, height = image.size
    ratio = width / max(1, height)
    draw_height = 8.5
    draw_width = draw_height * ratio
    if draw_width > 140:
        draw_width = 140
        draw_height = draw_width / ratio
    paragraph.add_run().add_picture(
        str(path),
        width=Mm(draw_width),
        height=Mm(draw_height),
    )


def _render_problem(
    document: Any,
    problem: dict[str, Any],
    number: int,
) -> None:
    title = _problem_title(problem, number)
    document.add_heading(title, level=3)
    model = content_blocks_for_problem(problem)
    blocks = model.get("blocks")
    if not isinstance(blocks, list):
        raise ValueError(f"\u9898\u76ee {problem.get('filename')} \u7f3a\u5c11\u5185\u5bb9\u5757")
    first_text = True
    paragraph = None
    for row in content_block_rows(model):
        if len(row) == 1 and row[0].get("type") == "image":
            asset = str(row[0].get("asset") or "")
            if not asset:
                raise ValueError(
                    f"\u9898\u76ee {problem.get('filename')} \u7684\u56fe\u7247\u5757\u7f3a\u5c11\u8d44\u6e90\u540d"
                )
            _add_image(document, _resource_path(problem, asset))
            paragraph = None
            continue
        standalone = len(row) == 1 and bool(row[0].get("display"))
        if paragraph is None or standalone:
            paragraph = document.add_paragraph()
        if standalone:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for block in row:
            block_type = block.get("type")
            if block_type == "text":
                text = str(block.get("text") or "")
                if first_text:
                    text = re.sub(
                        rf"^\s*{re.escape(title)}\s*",
                        "",
                        text,
                        count=1,
                    )
                    first_text = False
                if _UNSUPPORTED_MATH.search(text):
                    raise ValueError(
                        f"\u9898\u76ee {problem.get('filename')} \u542b\u6709\u65e0\u5750\u6807\u7684\u590d\u6742\u516c\u5f0f\uff0c\u8bf7\u91cd\u65b0\u5904\u7406\u540e\u5bfc\u51fa"
                    )
                if text:
                    _set_run_font(paragraph.add_run(text))
                continue
            if block_type != "latex":
                raise ValueError(
                    f"\u9898\u76ee {problem.get('filename')} \u5b58\u5728\u672a\u77e5\u5185\u5bb9\u5757 {block_type}"
                )
            source_text = str(block.get("source_text") or "")
            if not source_text and not block.get("formula_id"):
                raise ValueError(
                    f"\u9898\u76ee {problem.get('filename')} \u7684\u516c\u5f0f\u5757\u7f3a\u5c11\u539f\u6587\uff0c\u8bf7\u91cd\u65b0\u5904\u7406"
                )
            latex = str(block.get("latex") or "")
            fallback_asset = str(block.get("clean_crop_asset") or "")
            unresolved = block.get("recognition_state") in {
                "needs_review",
                "image_fallback",
                "human_verified_image",
            }
            if unresolved and fallback_asset:
                _add_inline_image(
                    paragraph,
                    _resource_path(problem, fallback_asset),
                )
                continue
            try:
                _append_omml(paragraph, latex)
            except (FormulaValidationError, ValueError):
                if not fallback_asset:
                    raise
                _add_inline_image(
                    paragraph,
                    _resource_path(problem, fallback_asset),
                )
        if standalone:
            paragraph = None
    answer_space = document.add_paragraph()
    answer_space.paragraph_format.space_after = Mm(45)


def export_docx(
    batch_id: str,
    problems: list[dict[str, Any]],
    output_dir: Path,
    *,
    allow_partial: bool = False,
) -> DocxExportResult:
    eligible = [problem for problem in problems if docx_exportable(problem)]
    skipped = [problem for problem in problems if not docx_exportable(problem)]
    if skipped and not allow_partial:
        first = skipped[0].get("filename") or skipped[0].get("id")
        raise ValueError(
            f"\u4ecd\u6709 {len(skipped)} \u9053\u9898\u672a\u901a\u8fc7\u786e\u8ba4\uff0c\u9996\u4e2a\u4e3a {first}"
        )
    if not eligible:
        raise ValueError("\u6ca1\u6709\u53ef\u5bfc\u51fa\u7684\u9898\u76ee")

    document = Document()
    _configure_document(document)
    title = document.add_heading("\u5c0f\u5965\u9519\u9898\u96c6", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    metadata = document.add_paragraph(
        f"\u751f\u6210\u65f6\u95f4\uff1a{datetime.now().astimezone().strftime('%Y-%m-%d %H:%M')}"
    )
    metadata.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_toc(document)
    document.add_page_break()

    grouped = _group_problems(eligible)
    problem_number = 0
    for group, categories in grouped.items():
        document.add_heading(group, level=1)
        for category, items in categories.items():
            document.add_heading(f"{category}\uff08{len(items)} \u9898\uff09", level=2)
            for problem in items:
                problem_number += 1
                _render_problem(document, problem, problem_number)
                if problem_number % 2 == 0 and problem_number < len(eligible):
                    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_batch = re.sub(r"[^A-Za-z0-9_-]", "", batch_id)[:8] or "batch"
    filename = f"mistake-book-{safe_batch}.docx"
    target = output_dir / f"{uuid.uuid4().hex}.docx"
    temporary = output_dir / f".{target.name}.{os.getpid()}.tmp"
    try:
        document.save(temporary)
        with zipfile.ZipFile(temporary) as archive:
            if "[Content_Types].xml" not in archive.namelist():
                raise ValueError("DOCX \u6587\u4ef6\u7ed3\u6784\u65e0\u6548")
        with temporary.open("rb") as handle:
            os.fsync(handle.fileno())
        os.replace(temporary, target)
    finally:
        temporary.unlink(missing_ok=True)
    return DocxExportResult(path=target, filename=filename)
